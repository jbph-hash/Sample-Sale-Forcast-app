import os
import random
import threading
import time
from datetime import datetime, timedelta
from math import ceil

from flask import Flask, render_template, request, redirect, url_for, jsonify, send_file
import sqlite3
import pandas as pd

# ---- Optional forecasting backends (best effort on Python 3.13) ----
USE_PM = False
USE_SM = False
try:
    import pmdarima as pm
    USE_PM = True
except Exception:
    try:
        from statsmodels.tsa.arima.model import ARIMA
        USE_SM = True
    except Exception:
        USE_PM = False
        USE_SM = False

from scipy.optimize import linprog
from openpyxl.chart import LineChart, Reference
from docx import Document
from docx.shared import Inches
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

APP_TITLE = "Retail POS + Forecasting + Inventory Optimization"
DB_FILE = "store.db"
REPORT_DIR = "reports"
os.makedirs(REPORT_DIR, exist_ok=True)

# Behavior toggles
AUTO_ADD_PRODUCTS_ON_SCAN = True  # default; actual runtime value stored in DB meta

# --------------------------
# Flask app
# --------------------------
app = Flask(__name__)

# --------------------------
# Database helpers
# --------------------------
def db_connect():
    return sqlite3.connect(DB_FILE, check_same_thread=False)

def init_db():
    with db_connect() as conn:
        c = conn.cursor()
        c.execute("""
            CREATE TABLE IF NOT EXISTS products (
                product_code TEXT PRIMARY KEY,
                product_name TEXT,
                price REAL,
                stock INTEGER DEFAULT 0,
                lead_time_days INTEGER DEFAULT 2,
                lead_time_minutes INTEGER DEFAULT 2880,
                holding_cost REAL DEFAULT 1.0,
                ordering_cost REAL DEFAULT 50.0,
                service_level REAL DEFAULT 0.95
            )
        """)
        c.execute("""
            CREATE TABLE IF NOT EXISTS transactions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                product_code TEXT,
                product_name TEXT,
                quantity INTEGER,
                price REAL,
                timestamp TEXT,
                receipt_id TEXT
            )
        """)
        c.execute("""
            CREATE TABLE IF NOT EXISTS meta (
                k TEXT PRIMARY KEY,
                v TEXT
            )
        """)
        c.execute("""
            CREATE TABLE IF NOT EXISTS reorders (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                product_code TEXT,
                product_name TEXT,
                suggested_qty INTEGER,
                created_at TEXT,
                status TEXT DEFAULT 'pending'
            )
        """)
        # Migrations for existing DBs
        try:
            c.execute("ALTER TABLE products ADD COLUMN lead_time_minutes INTEGER DEFAULT 2880")
        except Exception:
            pass
        try:
            c.execute("ALTER TABLE transactions ADD COLUMN receipt_id TEXT")
        except Exception:
            pass
        # seed default config if not present
        c.execute("INSERT OR IGNORE INTO meta(k, v) VALUES(?, ?)", ("AUTO_ADD_PRODUCTS_ON_SCAN", "true" if AUTO_ADD_PRODUCTS_ON_SCAN else "false"))
        c.execute("INSERT OR IGNORE INTO meta(k, v) VALUES(?, ?)", ("AUTO_REORDER_ENABLED", "true"))
        conn.commit()

init_db()

# --------------------------
# Utility functions
# --------------------------
def today_str():
    return datetime.now().strftime("%Y-%m-%d")

def date_only(ts_str):
    return ts_str.split(" ")[0]

def to_excel_multi(df_txn, df_stock, df_plan, filename):
    path = os.path.join(REPORT_DIR, filename)
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        if df_txn is not None and not df_txn.empty:
            df_txn.to_excel(writer, index=False, sheet_name="Transactions")
        if df_stock is not None and not df_stock.empty:
            df_stock.to_excel(writer, index=False, sheet_name="Stock")
        if df_plan is not None and not df_plan.empty:
            df_plan.to_excel(writer, index=False, sheet_name="Recommendations")
    return path

# --------------------------
# Config helpers
# --------------------------
def get_config_flag(flag_name: str, default_bool: bool) -> bool:
    try:
        with db_connect() as conn:
            c = conn.cursor()
            c.execute("SELECT v FROM meta WHERE k=?", (flag_name,))
            row = c.fetchone()
            if not row:
                return default_bool
            val = str(row[0]).strip().lower()
            return val in ("1", "true", "yes", "on")
    except Exception:
        return default_bool

def set_config_flag(flag_name: str, value_bool: bool) -> None:
    with db_connect() as conn:
        c = conn.cursor()
        c.execute("INSERT OR REPLACE INTO meta(k, v) VALUES(?, ?)", (flag_name, "true" if value_bool else "false"))
        conn.commit()

def is_auto_reorder_enabled() -> bool:
    return get_config_flag("AUTO_REORDER_ENABLED", True)

def maybe_create_reorder(conn, product_code: str, product_name: str, current_stock: int, rop: float, eoq_qty: float):
    if not is_auto_reorder_enabled():
        return
    if current_stock > rop:
        return
    qty = max(1, int(round(eoq_qty)))
    c = conn.cursor()
    # avoid duplicates pending for same product
    c.execute("SELECT COUNT(*) FROM reorders WHERE product_code=? AND status='pending'", (product_code,))
    if c.fetchone()[0] > 0:
        return
    c.execute(
        """
        INSERT INTO reorders (product_code, product_name, suggested_qty, created_at, status)
        VALUES (?, ?, ?, ?, 'pending')
        """,
        (product_code, product_name, qty, datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
    )
    conn.commit()

# --------------------------
# Forecasting (3-tier fallback)
# --------------------------
def daily_series(df, product_code):
    s = df[df["product_code"] == product_code].copy()
    if s.empty:
        return None
    s["timestamp"] = pd.to_datetime(s["timestamp"])
    series = s.set_index("timestamp")["quantity"].resample("D").sum().fillna(0)
    return series

def forecast_next_7(df_all, product_code):
    """
    Returns a pandas Series of length 7 with forecasted daily demand.
    Fallback order: pmdarima -> statsmodels ARIMA -> naive avg of last 7/14 days.
    """
    series = daily_series(df_all, product_code)
    if series is None or len(series) < 3:
        # Not enough data; use a small, pragmatic baseline instead of zeros
        start = pd.Timestamp.now().normalize() + timedelta(days=1)
        idx = pd.date_range(start, periods=7, freq="D")
        return pd.Series([1.0]*7, index=idx)

    # Prefer pmdarima on Python 3.13
    if USE_PM:
        try:
            model = pm.auto_arima(series, seasonal=False, suppress_warnings=True, error_action="ignore")
            fc = model.predict(n_periods=7)
            start = pd.Timestamp.now().normalize() + timedelta(days=1)
            idx = pd.date_range(start, periods=7, freq="D")
            return pd.Series(fc, index=idx).clip(lower=0)
        except Exception:
            pass

    if USE_SM:
        try:
            model = ARIMA(series, order=(1,1,1))
            res = model.fit()
            fc = res.forecast(steps=7)
            start = pd.Timestamp.now().normalize() + timedelta(days=1)
            fc.index = pd.date_range(start, periods=7, freq="D")
            return fc.clip(lower=0)
        except Exception:
            pass

    # Naive fallback: average of last 7 or last 14 days
    hist = series[-30:] if len(series) >= 30 else (series[-14:] if len(series) >= 14 else series)
    avg = float(hist.mean()) if len(hist) else 0.0
    start = pd.Timestamp.now().normalize() + timedelta(days=1)
    idx = pd.date_range(start, periods=7, freq="D")
    # Ensure a non-zero floor so we don't recommend zero by default
    floor = 1.0 if avg == 0.0 else avg
    return pd.Series([floor]*7, index=idx)

# --------------------------
# Optimization
# --------------------------
def eoq(QD, K, h):
    """
    EOQ: sqrt(2 * D * K / h)
    QD = demand over planning horizon (e.g., 30 days or 7 days sum)
    K  = ordering cost per order
    h  = holding cost per unit per period
    """
    try:
        from math import sqrt
        return max(0.0, (2.0 * QD * K / max(h, 1e-6)) ** 0.5)
    except Exception:
        return 0.0

def reorder_point(mu_daily, sigma_daily, L, z=1.65):
    """
    ROP = mu*L + z * sigma * sqrt(L)
    """
    from math import sqrt
    return max(0.0, mu_daily * L + z * sigma_daily * (sqrt(L)))

def optimize_week_plan(forecast_7, holding_cost=1.0, ordering_cost=50.0):
    """
    Simple LP: buy nonnegative x_t to cover demand across 7 days; minimize holding.
    Min sum(h * x_t) s.t. sum(x_t) >= sum(demand)
    This is a toy LP to demonstrate prescriptive analytics.
    """
    demand = list(float(x) for x in forecast_7)
    n = len(demand)
    c = [holding_cost] * n
    A = [[1.0]*n]
    b = [sum(demand)]
    bounds = [(0, None)] * n
    res = linprog(c, A_ub=A, b_ub=b, bounds=bounds, method="highs")
    if res.success:
        return [float(x) for x in res.x]
    # fallback: buy everything day 1
    buy = [0.0]*n
    buy[0] = sum(demand)
    return buy

def z_from_service_level(service_level: float) -> float:
    # Simple mapping for common service levels
    if service_level >= 0.99:
        return 2.33
    if service_level >= 0.975:
        return 1.96
    if service_level >= 0.95:
        return 1.65
    if service_level >= 0.9:
        return 1.28
    return 1.0

# --------------------------
# Routes / Views
# --------------------------
@app.route("/")
def index():
    with db_connect() as conn:
        df_tx = pd.read_sql_query(
            """
            SELECT t.*, p.stock AS current_stock
            FROM transactions t
            LEFT JOIN products p ON p.product_code = t.product_code
            ORDER BY t.timestamp DESC
            LIMIT 50
            """,
            conn,
        )
        df_products = pd.read_sql_query("SELECT * FROM products ORDER BY product_name", conn)
        # KPIs
        today = today_str()
        row_today = pd.read_sql_query(
            "SELECT COUNT(*) AS n, COALESCE(SUM(quantity*price),0) AS revenue FROM transactions WHERE DATE(timestamp)=?",
            conn,
            params=(today,),
        ).iloc[0]
        total_products = int(df_products.shape[0])
        total_stock_units = int(df_products["stock"].fillna(0).sum()) if not df_products.empty and "stock" in df_products.columns else 0
        kpis = {
            "total_products": total_products,
            "total_stock_units": total_stock_units,
            "today_transactions": int(row_today["n"] or 0),
            "today_revenue": float(row_today["revenue"] or 0.0),
        }
    return render_template(
        "index.html",
        title=APP_TITLE,
        rows=df_tx.to_dict(orient="records"),
        products=df_products.to_dict(orient="records"),
        today=today_str(),
        kpis=kpis,
        auto_add=get_config_flag("AUTO_ADD_PRODUCTS_ON_SCAN", AUTO_ADD_PRODUCTS_ON_SCAN)
    )

@app.route("/add_product", methods=["POST"])
def add_product():
    product_code = request.form["product_code"].strip()
    product_name = request.form["product_name"].strip()
    price = float(request.form.get("price", 0) or 0)
    stock = int(request.form.get("stock", 0) or 0)
    # Accept either minutes or days from form; prefer minutes
    lt_minutes = request.form.get("lead_time_minutes")
    lt_days = request.form.get("lead_time_days")
    lead_time_minutes = int(lt_minutes) if lt_minutes not in (None, "") else (int(lt_days) * 1440 if lt_days not in (None, "") else 2880)
    hc = float(request.form.get("holding_cost", 1.0) or 1.0)
    oc = float(request.form.get("ordering_cost", 50.0) or 50.0)
    sl = float(request.form.get("service_level", 0.95) or 0.95)

    with db_connect() as conn:
        c = conn.cursor()
        c.execute("""
            INSERT OR REPLACE INTO products
            (product_code, product_name, price, stock, lead_time_days, lead_time_minutes, holding_cost, ordering_cost, service_level)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (product_code, product_name, price, stock, ceil(lead_time_minutes/1440), lead_time_minutes, hc, oc, sl))
        conn.commit()
    return redirect(url_for("index"))

@app.route("/add_transaction", methods=["POST"])
def add_transaction():
    product_code = request.form["product_code"].strip()
    qty = int(request.form["quantity"] or 0)
    ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with db_connect() as conn:
        c = conn.cursor()
        c.execute("SELECT product_name, price, stock FROM products WHERE product_code=?", (product_code,))
        row = c.fetchone()
        if not row:
            if get_config_flag("AUTO_ADD_PRODUCTS_ON_SCAN", AUTO_ADD_PRODUCTS_ON_SCAN):
                # Create a placeholder product
                placeholder_name = product_code
                c.execute(
                    """
                    INSERT INTO products (product_code, product_name, price, stock, lead_time_days, lead_time_minutes, holding_cost, ordering_cost, service_level)
                    VALUES (?, ?, ?, ?, 2, 2880, 1.0, 50.0, 0.95)
                    """,
                    (product_code, placeholder_name, 0.0, 0),
                )
                conn.commit()
                c.execute("SELECT product_name, price, stock FROM products WHERE product_code=?", (product_code,))
                row = c.fetchone()
            else:
                return "Product not found. Add it first.", 400
        name, price, stock = row
        if stock < qty:
            # Distinguish out-of-stock vs just insufficient
            msg = "Out of stock." if stock <= 0 else "Insufficient stock."
            return msg, 400
        # Ensure there is an open receipt session
        c.execute("SELECT v FROM meta WHERE k='open_receipt_id'")
        r = c.fetchone()
        if not r:
            rid = f"R{int(time.time())}"
            c.execute("INSERT OR REPLACE INTO meta(k, v) VALUES('open_receipt_id', ?)", (rid,))
            c.execute("INSERT OR REPLACE INTO meta(k, v) VALUES('open_receipt_started_at', ?)", (ts,))
        else:
            rid = r[0]
        # reduce stock
        c.execute("UPDATE products SET stock = stock - ? WHERE product_code=?", (qty, product_code))
        # log txn
        c.execute("""
            INSERT INTO transactions (product_code, product_name, quantity, price, timestamp, receipt_id)
            VALUES (?, ?, ?, ?, ?, ?)
        """, (product_code, name, qty, price, ts, rid))
        # Auto-reorder check
        # Compute simple ROP/EOQ using current stats
        df_tx = pd.read_sql_query("SELECT * FROM transactions WHERE product_code=?", conn, params=(product_code,))
        series = daily_series(df_tx, product_code)
        mu = float(series.mean()) if series is not None and len(series) else 1.0
        sigma = float(series.std()) if series is not None and len(series) > 1 else max(0.5, 0.25 * mu)
        # Use minutes-based lead time
        c.execute("SELECT lead_time_minutes, holding_cost, ordering_cost, stock FROM products WHERE product_code=?", (product_code,))
        prow = c.fetchone()
        lt_minutes = int(prow[0]) if prow and prow[0] is not None else 2880
        holding = float(prow[1]) if prow and prow[1] is not None else 1.0
        ordering = float(prow[2]) if prow and prow[2] is not None else 50.0
        new_stock = int(prow[3]) if prow and prow[3] is not None else 0
        L_days = max(1, ceil(lt_minutes / 1440))
        z = 1.65
        rop_val = reorder_point(mu, sigma, L_days, z=z)
        eoq_qty = eoq(max(mu * 7, mu * 7), ordering, max(holding, 0.0001))
        maybe_create_reorder(conn, product_code, name, new_stock, rop_val, eoq_qty)
        conn.commit()
    return redirect(url_for("index"))

@app.route("/config", methods=["POST"])
def update_config():
    auto_add = request.form.get("auto_add_products") == "on"
    set_config_flag("AUTO_ADD_PRODUCTS_ON_SCAN", auto_add)
    return redirect(url_for("index"))

@app.route("/inventory")
def inventory_json():
    with db_connect() as conn:
        df = pd.read_sql_query("SELECT * FROM products ORDER BY product_name", conn)
    return jsonify(df.to_dict(orient="records"))

# --------------------------
# Reports
# --------------------------
def fetch_daily_transactions(date_str):
    with db_connect() as conn:
        q = """
        SELECT * FROM transactions
        WHERE DATE(timestamp) = ?
        ORDER BY timestamp
        """
        df = pd.read_sql_query(q, conn, params=(date_str,))
        df_prod = pd.read_sql_query("SELECT * FROM products", conn)
    return df, df_prod

def generate_daily_report_for(date_str):
    df_txn, df_stock = fetch_daily_transactions(date_str)
    if df_txn.empty:
        return None
    # create a placeholder recommendations sheet (optional)
    df_plan = pd.DataFrame({"note": [f"Daily report for {date_str}"]})
    filename = f"daily_report_{date_str}.xlsx"
    return to_excel_multi(df_txn, df_stock, df_plan, filename)

@app.route("/report/today")
def report_today():
    date_str = today_str()
    path = generate_daily_report_for(date_str)
    if not path:
        return f"No transactions on {date_str}."
    return send_file(path, as_attachment=True)

@app.route("/report/<date_str>")
def report_by_date(date_str):
    # date_str format: YYYY-MM-DD
    path = generate_daily_report_for(date_str)
    if not path:
        return f"No transactions on {date_str}."
    return send_file(path, as_attachment=True)

@app.route("/close_day", methods=["POST"])
def close_day():
    """Manual end-of-day: generate report and return link."""
    date_str = today_str()
    path = generate_daily_report_for(date_str)
    if not path:
        return jsonify({"status": "empty", "message": f"No transactions on {date_str}."})
    return jsonify({"status": "ok", "report": path})

# --------------------------
# Transactions utilities
# --------------------------
@app.route("/transactions/clear", methods=["POST"])
def transactions_clear():
    with db_connect() as conn:
        c = conn.cursor()
        c.execute("DELETE FROM transactions")
        conn.commit()
    return jsonify({"status": "ok"})

# --------------------------
# Analysis: Forecast + Optimization
# --------------------------
@app.route("/analyze")
def analyze():
    with db_connect() as conn:
        df_tx = pd.read_sql_query("SELECT * FROM transactions", conn)
        df_prod = pd.read_sql_query("SELECT * FROM products", conn)

    if df_tx.empty or df_prod.empty:
        return jsonify({"status": "error", "message": "Need transactions and products."})

    results = []
    for _, prod in df_prod.iterrows():
        code = prod["product_code"]
        name = prod["product_name"]
        holding = float(prod["holding_cost"])
        ordering = float(prod["ordering_cost"])
        lead_time_minutes = int(prod.get("lead_time_minutes", 2880) or 2880)
        lead_time = max(1, ceil(lead_time_minutes / 1440))
        service_level = float(prod["service_level"])
        current_stock = int(prod.get("stock", 0))

        fc7 = forecast_next_7(df_tx, code)  # pandas Series (7 days)
        plan = optimize_week_plan(fc7.values, holding_cost=holding, ordering_cost=ordering)

        # ROP & Safety stock (from history)
        series = daily_series(df_tx, code)
        mu = float(series.mean()) if series is not None and len(series) else 1.0  # baseline 1/day
        sigma = float(series.std()) if series is not None and len(series) > 1 else max(0.5, 0.25 * mu)
        z = z_from_service_level(service_level)

        rop = reorder_point(mu, sigma, max(1, lead_time), z=z)
        demand_7 = float(fc7.sum())
        # Ensure we plan for at least baseline demand in the week
        demand_7 = max(demand_7, 7.0 * mu)
        eoq_qty = eoq(max(demand_7, mu * 7), ordering, max(holding, 0.0001))

        # Simple order-up-to policy: cover lead time demand + safety stock vs current stock
        lead_time_demand = mu * max(1, lead_time)
        target_level = max(rop, lead_time_demand)
        suggested_order_qty = max(0.0, target_level - float(current_stock))

        # If EOQ suggests a larger, economical batch, take the max of the two
        suggested_order_qty = max(suggested_order_qty, eoq_qty if eoq_qty > 0 else 0.0)

        # Build a simple 7-day split: place the order on day 1 by default
        buy_split = [0.0] * 7
        if suggested_order_qty > 0:
            buy_split[0] = round(float(suggested_order_qty), 2)

        results.append({
            "product_code": code,
            "product_name": name,
            "forecast_7_total": round(demand_7, 2),
            "eoq_qty": round(eoq_qty, 2),
            "reorder_point": round(rop, 2),
            "current_stock": current_stock,
            "suggested_order_qty": round(float(suggested_order_qty), 2),
            "suggested_buy_split_7d": [round(x, 2) for x in buy_split]
        })

    df_plan = pd.DataFrame(results)
    # Save as a dated file
    fn = f"recommendations_{today_str()}.xlsx"
    path = os.path.join(REPORT_DIR, fn)
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        df_plan.to_excel(writer, index=False, sheet_name="Recommendations")

    return jsonify({"status": "ok", "saved_to": path, "results": results})

@app.route("/reorders")
def list_reorders():
    with db_connect() as conn:
        df = pd.read_sql_query("SELECT * FROM reorders ORDER BY created_at DESC", conn)
    return jsonify(df.to_dict(orient="records"))

@app.route("/reorders/mark", methods=["POST"])
def mark_reorder():
    rid = request.form.get("id")
    status = request.form.get("status", "ordered")
    if not rid:
        return "Missing id", 400
    with db_connect() as conn:
        c = conn.cursor()
        c.execute("UPDATE reorders SET status=? WHERE id=?", (status, rid))
        conn.commit()
    return redirect(url_for("index"))

@app.route("/ris/webhook", methods=["POST"])
def ris_webhook():
    data = request.get_json(silent=True) or {}
    code = str(data.get("product_code", "")).strip()
    delta = int(data.get("stock_delta", 0))
    if not code:
        return "Missing product_code", 400
    with db_connect() as conn:
        c = conn.cursor()
        c.execute("UPDATE products SET stock = COALESCE(stock,0) + ? WHERE product_code=?", (delta, code))
        conn.commit()
    return jsonify({"status": "ok"})

def _build_forecast_report():
    with db_connect() as conn:
        df_tx = pd.read_sql_query("SELECT * FROM transactions", conn)
        df_prod = pd.read_sql_query("SELECT * FROM products", conn)

        # Today's revenue
        today = today_str()
        row_today = pd.read_sql_query(
            "SELECT COUNT(*) AS n, COALESCE(SUM(quantity*price),0) AS revenue FROM transactions WHERE DATE(timestamp)=?",
            conn,
            params=(today,),
        ).iloc[0]

    if df_prod.empty:
        return None, "No products available."

    results = []
    per_product_forecasts = []
    # Accumulate total per day
    forecast_dates = pd.date_range(datetime.now()+timedelta(days=1), periods=7, freq="D")
    total_per_day = pd.Series([0.0]*7, index=forecast_dates)

    for _, prod in df_prod.iterrows():
        code = prod["product_code"]
        name = prod["product_name"]
        holding = float(prod["holding_cost"])
        ordering = float(prod["ordering_cost"])
        lead_time_minutes = int(prod.get("lead_time_minutes", 2880) or 2880)
        lead_time = max(1, ceil(lead_time_minutes / 1440))
        service_level = float(prod["service_level"])
        price = float(prod.get("price", 0.0))
        current_stock = int(prod.get("stock", 0))

        fc7 = forecast_next_7(df_tx, code)
        plan = optimize_week_plan(fc7.values, holding_cost=holding, ordering_cost=ordering)

        series = daily_series(df_tx, code)
        mu = float(series.mean()) if series is not None and len(series) else 1.0
        sigma = float(series.std()) if series is not None and len(series) > 1 else max(0.5, 0.25 * mu)
        z = z_from_service_level(service_level)

        rop = reorder_point(mu, sigma, max(1, lead_time), z=z)
        demand_7 = float(fc7.sum())
        demand_7 = max(demand_7, 7.0 * mu)
        eoq_qty = eoq(max(demand_7, mu * 7), ordering, max(holding, 0.0001))

        lead_time_demand = mu * max(1, lead_time)
        target_level = max(rop, lead_time_demand)
        suggested_order_qty = max(0.0, target_level - float(current_stock))
        suggested_order_qty = max(suggested_order_qty, eoq_qty if eoq_qty > 0 else 0.0)

        # Simple gains/loss estimates
        est_lost_sales_units = max(0.0, demand_7 - current_stock)
        est_lost_sales_value = est_lost_sales_units * price
        est_end_inventory_units = max(0.0, current_stock + suggested_order_qty - demand_7)
        est_holding_cost = est_end_inventory_units * holding
        est_net_impact = max(0.0, est_lost_sales_value - est_holding_cost)

        # Align by date string keys to avoid tz/clock drift KeyErrors
        fc_map = {pd.Timestamp(idx).normalize().strftime("%Y-%m-%d"): float(val) for idx, val in fc7.items()}
        row = {"product_code": code, "product_name": name}
        for d in forecast_dates:
            key = d.strftime("%Y-%m-%d")
            row[key] = fc_map.get(key, 0.0)
        row["total_7d"] = round(float(demand_7), 2)
        per_product_forecasts.append(row)
        total_per_day = total_per_day.add(fc7, fill_value=0)

        results.append({
            "product_code": code,
            "product_name": name,
            "price": price,
            "current_stock": current_stock,
            "forecast_7_total": round(demand_7, 2),
            "eoq_qty": round(eoq_qty, 2),
            "reorder_point": round(rop, 2),
            "suggested_order_qty": round(float(suggested_order_qty), 2),
            "est_lost_sales_value": round(float(est_lost_sales_value), 2),
            "est_holding_cost": round(float(est_holding_cost), 2),
            "est_net_impact": round(float(est_net_impact), 2),
        })

    # DataFrames
    df_forecast = pd.DataFrame(per_product_forecasts)
    totals_row = {"product_code": "TOTAL", "product_name": ""}
    for d in forecast_dates:
        key = d.strftime("%Y-%m-%d")
        totals_row[key] = round(float(total_per_day.loc[d]), 2)
    totals_row["total_7d"] = round(float(total_per_day.sum()), 2)
    df_forecast = pd.concat([df_forecast, pd.DataFrame([totals_row])], ignore_index=True)

    df_reco = pd.DataFrame(results)
    summary = {
        "generated_on": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "total_products": int(df_prod.shape[0]),
        "total_stock_units": int(df_prod.get("stock", pd.Series(dtype=int)).fillna(0).sum()) if not df_prod.empty else 0,
        "today_transactions": int(row_today["n"] or 0),
        "today_revenue": float(row_today["revenue"] or 0.0),
        "total_forecast_units_7d": round(float(total_per_day.sum()), 2),
        "sum_est_lost_sales_value": round(float(df_reco["est_lost_sales_value"].sum() if not df_reco.empty else 0.0), 2),
        "sum_est_holding_cost": round(float(df_reco["est_holding_cost"].sum() if not df_reco.empty else 0.0), 2),
        "sum_est_net_impact": round(float(df_reco["est_net_impact"].sum() if not df_reco.empty else 0.0), 2),
    }
    df_summary = pd.DataFrame([{k: v for k, v in summary.items()}])

    # Write Excel with chart
    fn = f"daily_forecast_{today_str()}.xlsx"
    path = os.path.join(REPORT_DIR, fn)
    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        df_summary.to_excel(writer, index=False, sheet_name="Summary")
        df_reco.to_excel(writer, index=False, sheet_name="Recommendations")
        df_forecast.to_excel(writer, index=False, sheet_name="Forecast")

        wb = writer.book
        ws = wb["Forecast"]
        # Determine chart range for the TOTAL row across the 7 date columns
        header = [cell.value for cell in ws[1]]
        # Date columns start after product_code and product_name => index 3 in 1-based excel indexing
        first_date_col = 3
        last_date_col = first_date_col + 7 - 1
        total_row_idx = ws.max_row  # last row is TOTAL

        # Create line chart of totals
        chart = LineChart()
        chart.title = "7-Day Total Forecast"
        data = Reference(ws, min_col=first_date_col, max_col=last_date_col, min_row=total_row_idx, max_row=total_row_idx)
        cats = Reference(ws, min_col=first_date_col, max_col=last_date_col, min_row=1, max_row=1)
        chart.add_data(data, titles_from_data=False)
        chart.set_categories(cats)
        chart.y_axis.title = "Units"
        chart.x_axis.title = "Date"
        ws.add_chart(chart, "B" + str(total_row_idx + 2))

    return path, None

@app.route("/analyze/download")
def analyze_download():
    path, err = _build_forecast_report()
    if err:
        return err, 400
    return send_file(path, as_attachment=True)

def _build_docx_forecast():
    path_xlsx, err = _build_forecast_report()
    if err:
        return None, err
    # Build 7-day total forecast series again for chart
    with db_connect() as conn:
        df_prod = pd.read_sql_query("SELECT * FROM products", conn)
        df_tx = pd.read_sql_query("SELECT * FROM transactions", conn)
    forecast_dates = pd.date_range(datetime.now().normalize()+timedelta(days=1), periods=7, freq='D')
    total_per_day = pd.Series([0.0]*7, index=forecast_dates)
    for _, prod in df_prod.iterrows():
        code = prod["product_code"]
        fc7 = forecast_next_7(df_tx, code)
        total_per_day = total_per_day.add(fc7, fill_value=0)

    # Plot chart to PNG
    fig, ax = plt.subplots(figsize=(6, 3))
    ax.plot([d.strftime('%Y-%m-%d') for d in forecast_dates], [float(total_per_day.loc[d]) for d in forecast_dates], marker='o')
    ax.set_title('7-Day Total Sales Forecast')
    ax.set_ylabel('Units')
    ax.set_xlabel('Date')
    ax.grid(True, alpha=0.3)
    img_path = os.path.join(REPORT_DIR, 'forecast_plot.png')
    fig.tight_layout()
    fig.savefig(img_path)
    plt.close(fig)

    # Create DOCX
    doc = Document()
    doc.add_heading('Daily Sales Forecast (7-Day)', level=1)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_picture(img_path, width=Inches(6.0))
    t = doc.add_table(rows=1, cols=2)
    hdr = t.rows[0].cells
    hdr[0].text = 'Date'
    hdr[1].text = 'Forecast Units'
    for d in forecast_dates:
        row = t.add_row().cells
        row[0].text = d.strftime('%Y-%m-%d')
        row[1].text = f"{total_per_day.loc[d]:.2f}"

    doc_fn = f"daily_forecast_{today_str()}.docx"
    doc_path = os.path.join(REPORT_DIR, doc_fn)
    doc.save(doc_path)
    return doc_path, None

@app.route("/analyze/download_docx")
def analyze_download_docx():
    path, err = _build_docx_forecast()
    if err:
        return err, 400
    return send_file(path, as_attachment=True)

# --------------------------
# Receipt finalize: compute lead time minutes and update products
# --------------------------
@app.route("/receipt/print", methods=["POST"])
def receipt_print():
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with db_connect() as conn:
        c = conn.cursor()
        c.execute("SELECT v FROM meta WHERE k='open_receipt_id'")
        rid_row = c.fetchone()
        c.execute("SELECT v FROM meta WHERE k='open_receipt_started_at'")
        start_row = c.fetchone()
        if not rid_row or not start_row:
            return jsonify({"status": "empty", "message": "No open receipt."})
        rid = rid_row[0]
        start_ts = datetime.strptime(start_row[0], "%Y-%m-%d %H:%M:%S")
        end_ts = datetime.strptime(now_str, "%Y-%m-%d %H:%M:%S")
        elapsed_min = max(1, int((end_ts - start_ts).total_seconds() // 60))
        # Affected products in this receipt
        c.execute("SELECT DISTINCT product_code FROM transactions WHERE receipt_id=?", (rid,))
        codes = [r[0] for r in c.fetchall()]
        # Update moving average of lead_time_minutes per product
        for code in codes:
            c.execute("SELECT lead_time_minutes FROM products WHERE product_code=?", (code,))
            row = c.fetchone()
            prev = int(row[0]) if row and row[0] is not None else elapsed_min
            updated = int(round(0.7 * prev + 0.3 * elapsed_min))
            c.execute("UPDATE products SET lead_time_minutes=?, lead_time_days=? WHERE product_code=?", (updated, max(1, ceil(updated/1440)), code))
        # Close receipt
        c.execute("DELETE FROM meta WHERE k IN ('open_receipt_id','open_receipt_started_at')")
        conn.commit()
    return jsonify({"status": "ok", "receipt_id": rid, "elapsed_minutes": elapsed_min, "affected_products": codes})

@app.route("/forecast/<product_code>")
def forecast_product(product_code):
    with db_connect() as conn:
        df_tx = pd.read_sql_query("SELECT * FROM transactions", conn)
    fc7 = forecast_next_7(df_tx, product_code)
    return jsonify({
        "product_code": product_code,
        "forecast": [{"date": str(idx.date()), "qty": float(val)} for idx, val in fc7.items()]
    })

# --------------------------
# Demo data seeding (optional)
# --------------------------
@app.route("/seed_demo")
def seed_demo():
    """Create a demo product and backfill 14 days of transactions for forecasting."""
    code = "P001"
    name = "Demo Product"
    price = 9.99
    initial_stock = 200
    with db_connect() as conn:
        c = conn.cursor()
        # Upsert product with healthy stock
        c.execute(
            """
            INSERT INTO products (product_code, product_name, price, stock, lead_time_days, holding_cost, ordering_cost, service_level)
            VALUES (?, ?, ?, ?, 2, 1.0, 50.0, 0.95)
            ON CONFLICT(product_code) DO UPDATE SET
                product_name=excluded.product_name,
                price=excluded.price,
                stock=excluded.stock
            """,
            (code, name, price, initial_stock),
        )

        # Insert transactions for the past 14 days with random quantities
        inserted = 0
        for d in range(14, 0, -1):
            day = (datetime.now() - timedelta(days=d)).strftime("%Y-%m-%d")
            qty = random.randint(1, 6)
            ts = f"{day} 12:00:00"
            c.execute(
                """
                INSERT INTO transactions (product_code, product_name, quantity, price, timestamp)
                VALUES (?, ?, ?, ?, ?)
                """,
                (code, name, qty, price, ts),
            )
            inserted += 1
        conn.commit()

    return jsonify({
        "status": "ok",
        "message": "Seeded demo product and 14 days of transactions.",
        "product_code": code,
        "transactions_inserted": inserted,
    })

# --------------------------
# Simple daily scheduler (optional)
# --------------------------
def _seconds_until_midnight():
    now = datetime.now()
    nxt = (now + timedelta(days=1)).replace(hour=0, minute=0, second=10, microsecond=0)
    return max(5, (nxt - now).seconds)

def daily_report_scheduler():
    while True:
        try:
            # Sleep until just after midnight
            time.sleep(_seconds_until_midnight())
            # Generate report for the previous date (yesterday)
            yday = (datetime.now() - timedelta(days=1)).strftime("%Y-%m-%d")
            generate_daily_report_for(yday)
        except Exception:
            # Keep scheduler alive even if report fails
            time.sleep(60)

# Start background scheduler thread
t = threading.Thread(target=daily_report_scheduler, daemon=True)
t.start()

# --------------------------
# Main
# --------------------------
if __name__ == "__main__":
    # Configure via environment variables for live/LAN deployments
    host = os.getenv("HOST", "0.0.0.0")  # 0.0.0.0 to listen on all interfaces
    port = int(os.getenv("PORT", "5000"))
    debug = os.getenv("DEBUG", "false").lower() in ("1", "true", "yes", "on")
    app.run(debug=debug, host=host, port=port)


